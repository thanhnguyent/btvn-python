from flask import Flask, render_template, Response, request, redirect
from db import edit_post1
from db import add_post, get_post
from datetime import datetime
from docx import Document
from docx.shared import Inches

app= Flask(__name__, template_folder='template')

post_base= get_post()
post = {}
a=0
for p in post_base:
    a+=1
    post[str(a)] = p
doc = Document()
@app.route("/resignation-letter")
def regis():
    return render_template("resignation-letter.html")
@app.route("/resignation-letter", methods=["POST"])
def regiss():
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    doc.add_heading('Thông tin bản thân...', level=2)
    para= doc.add_paragraph("Tên tôi là ")
    para.add_run(fullname).bold = True
    para.add_run(". Mục đích tôi viết bức thư này là ")
    para.add_run(reason).italic = True
    doc.save('btvn blog/cv1.docx')
    return redirect("/resignation-letter", code=302)


@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/",methods=["GET"])
def index():
    return render_template("blog.html", post = post)

@app.route("/", methods=["POST"])
def new_movie():
    now = datetime.now()
    dt_string = now.strftime("%Y-%m-%d %H:%M:%S")
    time= "Post at: "+str(dt_string)
    title= request.form["title"]
    content= request.form["content"]
    b= str(len(post)+1)
    post[b]= {
        "time":time,
        "title":title,
        "content": content
    } 
    add_post(time,title,content)
    return redirect("/", code=302)

@app.route("/post/<post_id>", methods=["GET"])
def detail(post_id):
    post1 = post.get(post_id)
    title= post1["title"]
    a= post_id
    return Response(render_template("edit.html",post1=post1, title=title, a=a), status=404, mimetype="text/html")

@app.route("/post/<post_id>", methods=["POST"])
def edit_post(post_id):
    now = datetime.now()
    dt_string = now.strftime("%Y-%m-%d %H:%M:%S")
    time= "Post at: "+str(dt_string)
    title= request.form["title"]
    content= request.form["content"]
    b= str(post_id)
    post[str(post_id)]= {
        "time" : time,
        "title":title,
        "content": content
    } 
    edit_post1(b,time, title, content)
    return redirect("/", code=302)

if __name__== "__main__":
    app.run(debug=True)
